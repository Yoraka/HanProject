from ..indexing import UnihanIRGs
import pickle
import random
import os
import docx
def load_index_file():
    with open('.\output\index\index_file.pkl', 'rb') as file:
        index = pickle.load(file)
        # 打印index大小
        print(f'index:{len(index)}')
    return index
# 通过引入Unihan_IRGSources数据，与index_file.pkl结合，将汉字按照部首进行分类
def get_unicode_radical_index(index):
    unihan = UnihanIRGs.UnihanIRGs()
    unicode_radical_index = {}
    i = 0
    for unicode in index:
        try:
            unicode_hex = hex(ord(unicode))
        except:
            print(f'hex except:{unicode}')
            continue
        radical = unihan.query_unihan_irg_sources(f'U+{unicode_hex[2:].upper()}', 'kRSUnicode')
        if radical is not None:
            if unicode not in unicode_radical_index:
                unicode_radical_index[unicode] = []
            radical = radical.split('.')[0]
            # 检查radical是否为数字
            if radical.isdigit():
                pass
            else:
                #去掉最后一个字符
                radical = radical[:-1]
                if not radical.isdigit():
                    print(f'radical is not digit:{unicode} {radical}')
            unicode_radical_index[unicode].append(radical)
        else:
            print(f'radical is None:{unicode}')
        i += 1
    print(f'i:{i}')
    return unicode_radical_index

# 对每个部首，将其下的汉字的index value取出，然后返回value list
def get_unicodes_values(index, unicode_list):
    unicodes_values = {}
    for unicode in unicode_list:
        if unicode in index:
            value = index.get(unicode)
            unicodes_values[unicode] = value
        else:
            print(f'unicode not in index:{unicode}')
    return unicodes_values
def rebuild_entry(value):
    # 将输入的value按换行符分割，以便获取标题和内容
    lines = value.split('\n')
    if not lines:
        return ""

    # 第一行应该是标题
    title = lines[0].strip()
    # 从第二行开始是具体内容
    content = '\n'.join(lines[1:])
    # 重建条目格式
    rebuilt_entry = f"{title}##{title}\n{content}\n"

    return rebuilt_entry

#康熙部首Number从1到214，对应的unicode值从U+2F00到U+2FD5
def kxNumber2unicode(kxNumber):
    kxNumber = int(kxNumber)
    if kxNumber < 1 or kxNumber > 214:
        return None
    unicode = hex(0x2F00 + kxNumber - 1)
    return unicode

def run():
    index = load_index_file()
    unicode_radical_index = get_unicode_radical_index(index)
    print(f'unicode_radical_index:{len(unicode_radical_index)}')
    # value的值应该从1到214不等，检查这些值是否都存在
    values = []
    for key in unicode_radical_index:
        for value in unicode_radical_index[key]:
            if int(value) > 214:
                print(f'error:{key} {value}')
            else:
                values.append(value)
    values = list(set(values))
    print(len(values))
    # 按照部首进行分类，每个部首下有多个unicode汉字
    radical_unicode_index = {}
    for key in unicode_radical_index:
        for value in unicode_radical_index[key]:
            if value not in radical_unicode_index:
                radical_unicode_index[value] = []
            radical_unicode_index[value].append(key)
    # 将radical_unicode_index的key下的所有value作为list传给get_radical_hans_indexvalue
    n = 0
    random_n = random.randint(0, 213)
    # 检查.data/radicals_docx目录是否存在，存在则删除目录下所有文件，不存在则创建目录
    if os.path.exists('./data/radicals_docx'):
        for root, dirs, files in os.walk('./data/radicals_docx'):
            for file in files:
                os.remove(os.path.join(root, file))
    else:
        os.makedirs('./data/radicals_docx')
    for key in radical_unicode_index:
        unicode_values = get_unicodes_values(index, radical_unicode_index[key])
        #print(f'unicode_values_len:{len(unicode_values)}')
        # 重建entry
        rebuilt_entries = []
        for unicode in unicode_values:
            rebuilt_entries.append(rebuild_entry(unicode_values[unicode]))
        # 将rebuilt_entries写入docx，entries之间用换行符连接，文件名是部首的值，文件存入radicals_docx目录下，如果没有该目录则创建
        # 部首的值是康熙部首的值，按照utf8KX映射表，将康熙部首的值转换为unicode的值，然后转为汉字拼入文件名
        kxUnicode = kxNumber2unicode(key)
        isFiveRadical = None
        #如果 key == 61 or  149 or 30 or 76 or 9 则文件名前加0
        if key in ['61', '149', '30', '76', '9']:
            isFiveRadical = 'a'
        else:
            isFiveRadical = ''
        file_path = f'./data/radicals_docx/{isFiveRadical}{key}_{chr(int(kxUnicode, 16))}部.docx'
        # 检查文件是否存在, 不存在就创建, 用docx.Document()创建并写入
        os.makedirs(os.path.dirname(file_path), exist_ok=True)
        # 不需要检查是否已经存在, 因为每次都是重新写入
        # 写入字体使用等线中文五号
        doc = docx.Document()
        doc.styles['Normal'].font.name = u'等线'
        doc.styles['Normal']._element.rPr.rFonts.set(docx.oxml.ns.qn('w:eastAsia'), u'等线')
        doc.styles['Normal'].font.size = docx.shared.Pt(10.5)
        doc.styles['Normal'].font.bold = False
        doc.styles['Normal'].font.italic = False
        doc.styles['Normal'].font.underline = False
        doc.styles['Normal'].font.strike = False
        doc.styles['Normal'].font.all_caps = False
        
        for entry in rebuilt_entries:
            doc.add_paragraph(entry)
        doc.save(file_path)

    return 1

if __name__ == '__main__':
    run()